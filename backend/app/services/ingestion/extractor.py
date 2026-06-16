import csv
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class UnsupportedDocumentTypeError(ValueError):
    pass


class TextExtractionError(RuntimeError):
    pass


class TextExtractor:
    """Extract text from supported enterprise document formats.

    The extractor never performs OCR. Scanned PDFs will therefore return little or no text;
    OCR should be added as a separate, explicit pipeline step when needed.
    """

    supported_extensions = {".txt", ".md", ".csv", ".pdf", ".docx"}

    def extract(self, path: Path) -> list[tuple[int | None, str]]:
        suffix = path.suffix.lower()
        if suffix not in self.supported_extensions:
            raise UnsupportedDocumentTypeError(f"Unsupported file type: {suffix}")

        try:
            if suffix in {".txt", ".md"}:
                return [(None, path.read_text(encoding="utf-8", errors="ignore"))]
            if suffix == ".csv":
                return [(None, self._extract_csv(path))]
            if suffix == ".pdf":
                return self._extract_pdf(path)
            if suffix == ".docx":
                return self._extract_docx(path)
        except Exception as exc:
            raise TextExtractionError(f"Could not extract text from {path.name}: {exc}") from exc

    def _extract_pdf(self, path: Path) -> list[tuple[int | None, str]]:
        reader = PdfReader(str(path))
        pages: list[tuple[int | None, str]] = []
        for index, page in enumerate(reader.pages, start=1):
            pages.append((index, page.extract_text() or ""))
        return pages

    def _extract_docx(self, path: Path) -> list[tuple[int | None, str]]:
        doc = DocxDocument(str(path))
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return [(None, "\n".join(parts))]

    def _extract_csv(self, path: Path) -> str:
        rows: list[str] = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as file:
            sample = file.read(4096)
            file.seek(0)
            dialect = csv.Sniffer().sniff(sample) if sample.strip() else csv.excel
            reader = csv.reader(file, dialect)
            for row in reader:
                rows.append(" | ".join(cell.strip() for cell in row))
        return "\n".join(rows)
